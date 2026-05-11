"""
Docx Parser - 解析 Word (`.docx`) 檔案中的題目
使用 python-docx 解析 Word 表格，抽取題目資訊和圖片
"""
import re
import os
from pathlib import Path
from typing import List, Optional, Tuple
from docx import Document
from docx.oxml import parse_xml

from .models import Question


class DocxParser:
    """解析 docx 檔案中的題目"""

    def __init__(self, docx_path: str, image_dir: str):
        """
        初始化解析器

        Args:
            docx_path: docx 檔案路徑
            image_dir: 圖片輸出目錄
        """
        self.docx_path = docx_path
        self.image_dir = Path(image_dir)
        self.image_dir.mkdir(parents=True, exist_ok=True)

        # 從檔名提取年份和科目
        self.year, self.subject = self._parse_filename()

        self.doc = Document(docx_path)
        self.questions: List[Question] = []

    def _parse_filename(self) -> Tuple[str, str]:
        """
        從檔名解析年份和科目

        檔名格式: {year}-{subject}-考題分類.docx
        例如: 105-國文-考題分類.docx -> year=105, subject=國文
        """
        filename = Path(self.docx_path).stem  # 去除副檔名
        parts = filename.split("-")

        if len(parts) >= 2:
            year = parts[0]
            subject = parts[1]
        else:
            raise ValueError(f"無法解析檔名: {filename}")

        return year, subject

    def _extract_year_and_number(self, text: str) -> Tuple[str, int]:
        """
        從出處文字提取年份和題號

        格式: 出處：105 會考第 1 題 或 出處：105 英文會考第 1 題 或 出處：106英文會考第 4 題
        """
        pattern = r'出處：(\d+)\s*.*?會考第\s*(\d+)\s*題'
        match = re.search(pattern, text)

        if match:
            year = match.group(1)
            number = int(match.group(2))
            return year, number

        raise ValueError(f"無法從文字提取年份和題號: {text}")

    def _extract_category(self, text: str) -> Tuple[str, str, str]:
        """
        從分類文字提取 L1, L2, L3

        格式: 分類：【字詞語句與閱讀】白話閱讀－邏輯推理與文意閱讀 出處：105 會考第 1 題
        """
        # 先提取「出處」之前的分類部分
        category_match = re.search(r'分類：(.*?)出處：', text)
        if category_match:
            category_text = category_match.group(1).strip()
        else:
            category_text = text

        # 提取【...】中的內容作為第一層分類
        level_1_match = re.search(r'【(.*?)】', category_text)
        level_1 = level_1_match.group(1) if level_1_match else ""

        # 提取剩下的分類
        rest_text = category_text
        if level_1_match:
            rest_text = category_text[level_1_match.end():]

        # 用 － 分割獲取 L2 和 L3
        parts = rest_text.split('－')
        level_2 = parts[0].strip() if len(parts) > 0 else ""
        level_3 = parts[1].strip() if len(parts) > 1 else ""

        return level_1, level_2, level_3

    def _parse_question_content(self, text: str) -> dict:
        """
        解析題目內容，提取題幹和選項

        格式:
        分類：【...】... 出處：... 題目內容：題幹文字
        (A) 選項 A
        (B) 選項 B
        (C) 選項 C
        (D) 選項 D
        """
        # 先提取「題目內容：」之後的內容
        content_match = re.search(r'題目內容：(.*)', text, re.DOTALL)
        if content_match:
            text = content_match.group(1).strip()

        # 尋找 (A) (B) (C) (D) 的位置
        option_pattern = r'\n?\s*\(([ABCD])\)'
        matches = list(re.finditer(option_pattern, text))

        options = {"A": "", "B": "", "C": "", "D": ""}
        question_stem = text

        if matches:
            # 題幹是到第一個選項之前的內容
            question_stem = text[:matches[0].start()].strip()

            # 提取各個選項
            for i, match in enumerate(matches):
                option_letter = match.group(1)
                start = match.end()
                end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
                option_text = text[start:end].strip()
                options[option_letter] = option_text

        return {
            "question_stem": question_stem,
            "option_a": options.get("A", ""),
            "option_b": options.get("B", ""),
            "option_c": options.get("C", ""),
            "option_d": options.get("D", ""),
        }

    def _parse_explanation(self, text: str) -> dict:
        """
        解析試題解析，提取主解析和各選項解析
        """
        # 先檢查是否包含嵌套選項解析格式
        if '. ' in text and ('(A)' in text or '(B)' in text or '(C)' in text or '(D)' in text):
            # 使用嵌套選項解析的模式
            nested_start_match = re.search(r'\([ABCD]\)\]', text)
            if nested_start_match:
                # 提取主解析
                main_explanation = text[:nested_start_match.start()].strip()
                nested_text = text[nested_start_match.end():]

                # 提取各個選項解析
                option_pattern = r'\(([ABCD])\)'
                nested_matches = list(re.finditer(option_pattern, nested_text))

                explanations = {"A": "", "B": "", "C": "", "D": ""}
                for nested_match in nested_matches:
                    option_letter = nested_match.group(1)
                    start = nested_match.end()
                    end = nested_matches[nested_matches.index(nested_match) + 1].start() if nested_matches.index(nested_match) + 1 < len(nested_matches) else len(nested_text)
                    explanation_text = nested_text[start:end].strip()
                    explanations[option_letter] = explanation_text

                return {
                    "explanation": main_explanation,
                    "explanation_a": explanations.get("A", ""),
                    "explanation_b": explanations.get("B", ""),
                    "explanation_c": explanations.get("C", ""),
                    "explanation_d": explanations.get("D", ""),
                }

        # 一般格式處理
        option_pattern = r'\s*\(([ABCD])\)'
        matches = list(re.finditer(option_pattern, text))

        explanations = {"A": "", "B": "", "C": "", "D": ""}
        main_explanation = text

        if matches:
            main_explanation = text[:matches[0].start()].strip()

            for i, match in enumerate(matches):
                option_letter = match.group(1)
                start = match.end()
                end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
                explanation_text = text[start:end].strip()
                explanations[option_letter] = explanation_text

        return {
            "explanation": main_explanation,
            "explanation_a": explanations.get("A", ""),
            "explanation_b": explanations.get("B", ""),
            "explanation_c": explanations.get("C", ""),
            "explanation_d": explanations.get("D", ""),
        }

    def _extract_answer(self, text: str) -> str:
        """
        從答案文字提取正確答案

        格式: 正確答案：【D】
        """
        pattern = r'正確答案：【([ABCD])】'
        match = re.search(pattern, text)
        if match:
            return match.group(1)
        return ""

    def parse(self) -> List[Question]:
        """
        解析 docx 檔案，返回題目列表

        Returns:
            題目列表
        """
        questions = []

        for table in self.doc.tables:
            try:
                # 每個表格代表一個題目
                if len(table.rows) < 2:
                    continue

                # Row 0: 分類、出處、題目內容
                row0_text = table.rows[0].cells[0].text

                # Row 1: 正確答案、試題解析
                row1_text = table.rows[1].cells[0].text

                # 提取年份和題號
                year, number = self._extract_year_and_number(row0_text)

                # 提取分類
                level_1, level_2, level_3 = self._extract_category(row0_text)

                # 解析題目內容
                content = self._parse_question_content(row0_text)

                # 提取答案
                answer = self._extract_answer(row1_text)

                # 解析試題解析
                explanation_text = row1_text.split("試題解析：")[-1] if "試題解析：" in row1_text else ""
                explanation_data = self._parse_explanation(explanation_text)

                # 檢查是否有圖片
                has_image, image_filename = self._check_for_image(table)

                # 建立 Question 物件
                question = Question(
                    year=year,
                    number=number,
                    subject_category=self.subject,
                    level_1=level_1,
                    level_2=level_2,
                    level_3=level_3,
                    question_stem=content["question_stem"],
                    option_a=content["option_a"],
                    option_b=content["option_b"],
                    option_c=content["option_c"],
                    option_d=content["option_d"],
                    answer=answer,
                    explanation=explanation_data["explanation"],
                    explanation_a=explanation_data["explanation_a"],
                    explanation_b=explanation_data["explanation_b"],
                    explanation_c=explanation_data["explanation_c"],
                    explanation_d=explanation_data["explanation_d"],
                    has_question_image=has_image,
                    image_filename=image_filename,
                    category="圖片" if has_image else "純文字",
                    raw_category=row0_text.split("出處：")[0].replace("分類：", "").strip() if "出處：" in row0_text else "",
                )

                questions.append(question)

            except Exception as e:
                print(f"解析表格時發生錯誤: {e}")
                continue

        return questions

    def _check_for_image(self, table) -> Tuple[bool, Optional[str]]:
        """
        檢查表格中是否有圖片

        Returns:
            (has_image, image_filename)
        """
        # 檢查表格中的所有段落是否有圖片
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run._element.xpath('.//a:blip'):
                            # 找到圖片，生成檔名
                            image_filename = f"{self.year}-{self.subject}-Q{len(self.questions) + 1}"
                            return True, image_filename
        return False, None
